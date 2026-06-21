"""Convert contract-cnc-erp.docx to PDF using python-docx + reportlab.

Approach: two-pass build to know total pages for footer.
"""
import copy
import re
import sys
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from io import BytesIO

SRC = Path(r"E:\claude\smart-workshop-erp\docs\contract-cnc-erp.docx")
DST = Path(r"E:\claude\smart-workshop-erp\docs\contract-cnc-erp.pdf")


# ============ Font ============
def register_font():
    candidates = [
        (r"C:\Windows\Fonts\msyh.ttc", "Microsoft YaHei"),
        (r"C:\Windows\Fonts\msyh.ttf", "Microsoft YaHei"),
        (r"C:\Windows\Fonts\msyhbd.ttc", "Microsoft YaHei Bold"),
        (r"C:\Windows\Fonts\simhei.ttf", "SimHei"),
        (r"C:\Windows\Fonts\simsun.ttc", "SimSun"),
        (r"C:\Windows\Fonts\Deng.ttf", "DengXian"),
        ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
        ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "WenQuanYi"),
    ]
    for path, name in candidates:
        if Path(path).exists():
            try:
                pdfmetrics.registerFont(TTFont('CJK', path))
                print(f"  font loaded: {name} from {path}")
                return 'CJK'
            except Exception as e:
                continue
    try:
        pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
        print("  font fallback: STSong-Light (CID)")
        return 'STSong-Light'
    except Exception as e:
        print(f"  CID font failed: {e}")
        return 'Helvetica'


FONT = register_font()


# ============ Styles ============
def make_styles():
    s_title = ParagraphStyle('Title', fontName=FONT, fontSize=22, leading=30,
                             alignment=TA_CENTER, textColor=colors.HexColor('#1f2328'),
                             spaceBefore=12, spaceAfter=18)
    s_h1 = ParagraphStyle('H1', fontName=FONT, fontSize=15, leading=22,
                         spaceBefore=14, spaceAfter=8,
                         textColor=colors.HexColor('#0969da'),
                         keepWithNext=True)
    s_article = ParagraphStyle('Article', fontName=FONT, fontSize=14, leading=22,
                               spaceBefore=14, spaceAfter=6,
                               textColor=colors.HexColor('#1f2328'),
                               keepWithNext=True)
    s_h3 = ParagraphStyle('H3', fontName=FONT, fontSize=11, leading=18,
                         spaceBefore=8, spaceAfter=4,
                         textColor=colors.HexColor('#1f2328'),
                         keepWithNext=True)
    s_body = ParagraphStyle('Body', fontName=FONT, fontSize=10, leading=17,
                            alignment=TA_JUSTIFY, firstLineIndent=20,
                            textColor=colors.HexColor('#1f2328'),
                            spaceBefore=2, spaceAfter=2)
    s_center = ParagraphStyle('Center', fontName=FONT, fontSize=10, leading=17,
                             alignment=TA_CENTER, textColor=colors.HexColor('#1f2328'),
                             spaceBefore=2, spaceAfter=2)
    s_bullet = ParagraphStyle('Bullet', fontName=FONT, fontSize=10, leading=16,
                             leftIndent=24, bulletIndent=10, spaceBefore=0, spaceAfter=0)
    s_quote = ParagraphStyle('Quote', fontName=FONT, fontSize=9.5, leading=15,
                            leftIndent=24, rightIndent=24,
                            textColor=colors.HexColor('#666666'),
                            spaceBefore=4, spaceAfter=4)
    s_note = ParagraphStyle('Note', fontName=FONT, fontSize=9, leading=14,
                           alignment=TA_CENTER, textColor=colors.HexColor('#666666'),
                           spaceBefore=4, spaceAfter=4)
    return {
        'title': s_title, 'h1': s_h1, 'article': s_article, 'h3': s_h3,
        'body': s_body, 'center': s_center, 'bullet': s_bullet,
        'quote': s_quote, 'note': s_note,
    }


# ============ Inline formatting ============
def runs_to_markup(runs):
    parts = []
    for run in runs:
        text = run.text
        if not text:
            continue
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        text = text.replace('\n', '<br/>')
        if run.font.bold:
            text = f'<b>{text}</b>'
        if run.font.italic:
            text = f'<i>{text}</i>'
        if run.font.color and run.font.color.rgb:
            c = str(run.font.color.rgb)
            if c != '000000':
                text = f'<font color="#{c}">{text}</font>'
        parts.append(text)
    return ''.join(parts)


# ============ Table ============
def render_table(table):
    nrows = len(table.rows)
    ncols = len(table.columns) if table.rows else 0
    if nrows == 0 or ncols == 0:
        return None

    data = []
    for i, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            cell_paras = []
            for p in cell.paragraphs:
                markup = runs_to_markup(p.runs)
                if markup.strip():
                    cell_paras.append(markup)
            cell_text = '<br/>'.join(cell_paras) if cell_paras else ''
            if i == 0:
                # Header: bold
                cell_text = f'<b>{cell_text}</b>' if cell_text else ''
            style = ParagraphStyle(
                f'Cell_{i}', fontName=FONT, fontSize=9, leading=13,
                textColor=colors.HexColor('#1f2328'),
            )
            row_data.append(Paragraph(cell_text, style))
        data.append(row_data)

    page_width = A4[0] - 4 * cm
    col_widths = [page_width / ncols] * ncols
    if ncols == 2:
        col_widths = [page_width * 0.35, page_width * 0.65]
    elif ncols == 3:
        col_widths = [page_width * 0.16, page_width * 0.50, page_width * 0.34]
    elif ncols == 4:
        col_widths = [page_width * 0.10, page_width * 0.55, page_width * 0.15, page_width * 0.20]

    t = Table(data, colWidths=col_widths, repeatRows=1)
    style_cmds = [
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#888888')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9E2F3')),
    ]
    t.setStyle(TableStyle(style_cmds))
    return t


# ============ Body parser ============
def docx_to_flowables():
    doc = Document(SRC)
    styles = make_styles()
    flow = []

    body = doc.element.body
    paragraphs = doc.paragraphs
    tables = doc.tables
    para_idx = 0
    table_idx = 0

    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            if para_idx < len(paragraphs):
                render_paragraph(paragraphs[para_idx], flow, styles)
                para_idx += 1
        elif tag == 'tbl':
            if table_idx < len(tables):
                t = render_table(tables[table_idx])
                if t:
                    flow.append(t)
                    flow.append(Spacer(1, 6))
                table_idx += 1

    return flow


def render_paragraph(p, flow, styles):
    text = p.text.strip()

    # Empty
    if not text:
        flow.append(Spacer(1, 4))
        return

    # Main title
    if text == '技术开发（委托）合同':
        flow.append(Paragraph(text, styles['title']))
        return

    # Detect heading type
    is_article = bool(re.match(r'^第[一二三四五六七八九十]+条\s', text))
    is_appendix = text.startswith('附件') and len(text) < 30
    is_signature_block = text in ('甲方（委托方）', '乙方（受托方）', '双方签章', '签章页')
    is_subtitle = text.startswith('合同编号：') or text.startswith('签订日期：') or text.startswith('签订地点：') or text.startswith('本《技术开发')
    is_decoration = text in ('委托方（以下简称"甲方"）：昆山佰泰胜精密机械有限公司',
                             '受托方（以下简称"乙方"）：河南晓评信息科技有限公司')

    if is_signature_block:
        flow.append(Paragraph(text, styles['h1']))
        return
    if is_article:
        flow.append(Paragraph(text, styles['article']))
        return
    if is_appendix:
        flow.append(Paragraph(text, styles['h1']))
        return
    if is_subtitle or is_decoration:
        flow.append(Paragraph(text, styles['center']))
        return
    if text.startswith('● '):
        b = Paragraph(runs_to_markup(p.runs), styles['bullet'])
        flow.append(b)
        return
    if text.startswith('「') and text.endswith('」'):
        flow.append(Paragraph(text, styles['quote']))
        return
    if text.startswith('【注】'):
        flow.append(Paragraph(text, styles['quote']))
        return
    if text.startswith('⚠'):
        flow.append(Paragraph(text, styles['note']))
        return

    # Regular body with first-line indent and inline bold
    markup = runs_to_markup(p.runs)
    if markup.strip():
        flow.append(Paragraph(markup, styles['body']))


# ============ Header / Footer ============
TOTAL_PAGES = [0]


def make_canvas(total):
    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setFont(FONT, 9)
        canvas.setFillColor(colors.HexColor('#888888'))
        canvas.drawCentredString(
            A4[0] / 2, A4[1] - 1.2 * cm,
            '昆山佰泰胜专属 ERP 系统定制开发项目 · 技术开发（委托）合同'
        )
        canvas.drawCentredString(
            A4[0] / 2, 1.0 * cm,
            f'第 {doc.page} 页 / 共 {total} 页'
        )
        canvas.restoreState()
    return on_page


# ============ Main ============
def main():
    print(f"[INFO] Reading: {SRC}")

    # Build twice to know total pages
    flow = docx_to_flowables()
    print(f"[INFO] {len(flow)} flowables generated")

    # First pass: real build to file (count pages).
    # NOTE: reportlab mutates flowables during build (Paragraph computes wrap once,
    # Table splits rows, etc.), so we deep-copy for the second pass to avoid
    # producing an empty PDF.
    from reportlab.platypus import PageTemplate, BaseDocTemplate, Frame
    tmp_path = DST.with_suffix('.tmp.pdf')
    doc_tmp = SimpleDocTemplate(
        str(tmp_path), pagesize=A4,
        leftMargin=3 * cm, rightMargin=3 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )
    counter = {'page': 0}

    def counting_on_page(canvas, doc):
        counter['page'] = max(counter['page'], doc.page)

    flow_for_count = [copy.deepcopy(f) for f in flow]
    doc_tmp.build(flow_for_count, onFirstPage=counting_on_page, onLaterPages=counting_on_page)
    total = max(counter['page'], 1)
    print(f"[INFO] Total pages: {total}")

    # Remove temp file
    try:
        tmp_path.unlink()
    except FileNotFoundError:
        pass

    # Second pass: real PDF with footer
    flow_for_real = [copy.deepcopy(f) for f in flow]
    doc = SimpleDocTemplate(
        str(DST), pagesize=A4,
        leftMargin=3 * cm, rightMargin=3 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title='昆山佰泰胜专属 ERP 系统定制开发项目 · 技术开发（委托）合同',
        author='河南晓评信息科技有限公司',
    )
    doc.build(flow_for_real, onFirstPage=make_canvas(total), onLaterPages=make_canvas(total))

    size = DST.stat().st_size
    sys.stdout.reconfigure(encoding='utf-8')
    print(f"[OK] PDF: {DST}")
    print(f"     Size: {size:,} bytes")
    print(f"     Pages: {total}")


if __name__ == '__main__':
    main()
