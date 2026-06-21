"""Convert appendix md files (A/B/C/D) to standalone docx files.

Usage: python md_to_appendix_docx.py <input.md> <output.docx> [header_text]
"""
import os
import re
import sys
import time
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------- helpers ----------------

def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_table(doc, rows, header_bg='D9E2F3'):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            add_runs_with_bold(p, str(cell_text), font_name='宋体', size=10.5)
            if i == 0:
                for run in p.runs:
                    run.font.bold = True
                set_cell_bg(cell, header_bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
    return table


def set_para_font(para, font_name='宋体', size=10.5, bold=False, color=None):
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)


def add_heading(doc, text, level=1):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_para_font(p, '黑体', 22, True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
    elif level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_para_font(p, '黑体', 16, True)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_para_font(p, '黑体', 14, True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_para_font(p, '黑体', 12, True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)


def add_para(doc, text, indent=True, bold=False, align=None):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    add_runs_with_bold(p, text, font_name='宋体', size=10.5, bold=bold)
    return p


def add_runs_with_bold(para, text, font_name='宋体', size=10.5, bold=False, color=None):
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = para.add_run(inner)
            run.font.bold = True
        else:
            run = para.add_run(part)
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = run.font.bold or bold
        if color and not run.font.bold:
            run.font.color.rgb = RGBColor.from_string(color)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '888888')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    doc.add_page_break()


def add_code_block(doc, text):
    """Render ``` code block ``` as monospaced preformatted paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.right_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')


# ---------------- parser / renderer ----------------

def parse_sections(md_text):
    lines = md_text.split('\n')
    sections = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            content = []
            i += 1
            while i < n:
                nxt = lines[i]
                if re.match(r'^#{1,4}\s+', nxt):
                    break
                content.append(nxt)
                i += 1
            sections.append({
                'level': level,
                'title': title,
                'content': '\n'.join(content).strip(),
            })
        else:
            i += 1
    return sections


def render_section(doc, sec):
    level = sec['level']
    title = sec['title']
    content = sec['content']

    if level == 1:
        add_heading(doc, title, level=1)
    elif level == 2:
        add_heading(doc, title, level=2)
    elif level == 3:
        add_heading(doc, title, level=3)
    else:
        add_heading(doc, title, level=3)

    if not content:
        return

    lines = content.split('\n')
    i = 0
    n = len(lines)
    in_code = False
    code_buf = []
    while i < n:
        line = lines[i].rstrip()

        # Code fence
        if line.startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
                i += 1
                continue
            else:
                # flush code
                add_code_block(doc, '\n'.join(code_buf))
                in_code = False
                code_buf = []
                i += 1
                continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Blockquote
        if line.startswith('>'):
            quote_lines = []
            while i < n and lines[i].rstrip().startswith('>'):
                quote_lines.append(re.sub(r'^>\s?', '', lines[i].rstrip()))
                i += 1
            quote_text = ' '.join(q.strip() for q in quote_lines if q.strip())
            if quote_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.74)
                p.paragraph_format.right_indent = Cm(0.74)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if quote_text.startswith('【注】'):
                    add_runs_with_bold(p, '【注】' + quote_text.replace('【注】', ''),
                                       font_name='楷体', size=10, color='666666')
                elif quote_text.startswith('⚠'):
                    add_runs_with_bold(p, '⚠ ' + quote_text.replace('⚠', ''),
                                       font_name='楷体', size=10, color='A04040')
                else:
                    add_runs_with_bold(p, '「' + quote_text + '」', font_name='楷体', size=10)
            continue

        # Table
        if line.startswith('|'):
            table_rows = []
            while i < n and lines[i].rstrip().startswith('|'):
                row_text = lines[i].rstrip()
                if re.match(r'^\|[\s\-|:]+\|$', row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
            if table_rows:
                add_table(doc, table_rows)
                doc.add_paragraph()
            continue

        # Horizontal rule
        if line.strip() == '---':
            add_hr(doc)
            i += 1
            continue

        # Heading within content
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            sub_level = len(m.group(1))
            sub_title = m.group(2).strip()
            add_heading(doc, sub_title, level=3 if sub_level <= 3 else 3)
            i += 1
            continue

        if line.strip() == '':
            i += 1
            continue

        # List item
        m_ul = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        m_ol = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if m_ul or m_ol:
            while i < n:
                nxt = lines[i].rstrip()
                if not nxt:
                    i += 1
                    continue
                m_ul2 = re.match(r'^(\s*)[-*+]\s+(.+)$', nxt)
                m_ol2 = re.match(r'^(\s*)(\d+)\.\s+(.+)$', nxt)
                if m_ul2:
                    indent_spaces = len(m_ul2.group(1))
                    bullet_text = m_ul2.group(2)
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.74 + 0.5 * (indent_spaces // 2))
                    p.paragraph_format.first_line_indent = Cm(-0.5)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    add_runs_with_bold(p, '● ' + bullet_text, font_name='宋体', size=10.5)
                    i += 1
                elif m_ol2:
                    indent_spaces = len(m_ol2.group(1))
                    num = m_ol2.group(2)
                    list_text = m_ol2.group(3)
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.74 + 0.5 * (indent_spaces // 2))
                    p.paragraph_format.first_line_indent = Cm(-0.5)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    add_runs_with_bold(p, f'{num}. ' + list_text, font_name='宋体', size=10.5)
                    i += 1
                else:
                    break
            continue

        # Plain text paragraph
        text = line.strip()
        para_lines = [text]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if (nxt == '' or nxt.startswith('#') or nxt.startswith('|') or
                    nxt.startswith('>') or nxt.strip() == '---' or
                    nxt.startswith('```') or
                    re.match(r'^\s*[-*+]\s+', nxt) or re.match(r'^\s*\d+\.\s+', nxt)):
                break
            para_lines.append(nxt)
            i += 1
        p_text = ' '.join(l.strip() for l in para_lines if l.strip())
        if p_text:
            add_para(doc, p_text, indent=True)


# ---------------- main ----------------

def convert(src_path, dst_path, header_text):
    src = Path(src_path)
    dst = Path(dst_path)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

        header = section.header
        hp = header.paragraphs[0]
        hp.text = header_text
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            run.font.size = Pt(9)
            run.font.name = '宋体'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), '宋体')

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run('第 ')
        run.font.size = Pt(9)
        run.font.name = '宋体'
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run = fp.add_run(' 页 / 共 ')
        run.font.size = Pt(9)
        run.font.name = '宋体'
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'NUMPAGES'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run = fp.add_run(' 页')
        run.font.size = Pt(9)
        run.font.name = '宋体'

    sections = parse_sections(src.read_text(encoding='utf-8'))
    for sec in sections:
        render_section(doc, sec)
        if sec.get('level') == 1:
            doc.add_paragraph()

    # Save with lock-resistant replace
    tmp_path = dst.with_suffix(dst.suffix + '.tmp')
    doc.save(str(tmp_path))
    replaced = False
    for attempt in range(3):
        try:
            if dst.exists():
                stale = dst.with_name(dst.stem + f'.stale{int(time.time())}{dst.suffix}')
                try:
                    dst.rename(stale)
                except PermissionError:
                    time.sleep(0.5)
                    continue
                else:
                    try:
                        stale.unlink()
                    except OSError:
                        pass
            os.replace(str(tmp_path), str(dst))
            replaced = True
            break
        except PermissionError:
            time.sleep(0.5)

    if not replaced:
        alt = dst.with_name(dst.stem + f'.{int(time.time())}{dst.suffix}')
        os.replace(str(tmp_path), str(alt))
        print(f"[WARN] saved to {alt.name} (locked dst)")
        dst = alt

    size = dst.stat().st_size
    print(f"[OK] {dst.name}: {size:,} bytes")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python md_to_appendix_docx.py <input.md> <output.docx> [header_text]")
        sys.exit(1)
    src = sys.argv[1]
    dst = sys.argv[2]
    header = sys.argv[3] if len(sys.argv) > 3 else '附件'
    convert(src, dst, header)
