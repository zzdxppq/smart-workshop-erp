"""Convert ops-manual-cutter.docx to A4 landscape PDF.

Reuses the rendering helpers from docx_to_pdf.py but with landscape pages.
"""
import copy
import re
import sys
from io import BytesIO
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import landscape, A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
)

SRC = Path(r'E:\claude\smart-workshop-erp\docs\ops-manual-cutter.docx')
DST = Path(r'E:\claude\smart-workshop-erp\docs\ops-manual-cutter.pdf')
PAGE_SIZE = landscape(A4)  # (29.7cm, 21cm)


# ------------ font ------------
def register_font():
    candidates = [
        (r'C:\Windows\Fonts\msyh.ttc', 'Microsoft YaHei'),
        (r'C:\Windows\Fonts\msyh.ttf', 'Microsoft YaHei'),
    ]
    for path, name in candidates:
        if Path(path).exists():
            try:
                pdfmetrics.registerFont(TTFont('CJK', path))
                return 'CJK'
            except Exception:
                continue
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    return 'STSong-Light'


FONT = register_font()


# ------------ styles ------------
def make_styles():
    return {
        'h1': ParagraphStyle('H1', fontName=FONT, fontSize=14, leading=20,
                            spaceBefore=10, spaceAfter=4,
                            textColor=colors.HexColor('#1f2328'),
                            keepWithNext=True),
        'h2': ParagraphStyle('H2', fontName=FONT, fontSize=12, leading=18,
                            spaceBefore=6, spaceAfter=2,
                            textColor=colors.HexColor('#1f2328'),
                            keepWithNext=True),
        'body': ParagraphStyle('Body', fontName=FONT, fontSize=9, leading=14,
                              spaceBefore=1, spaceAfter=1),
        'quote': ParagraphStyle('Quote', fontName=FONT, fontSize=8.5, leading=13,
                                leftIndent=18, rightIndent=18,
                                textColor=colors.HexColor('#666666'),
                                spaceBefore=2, spaceAfter=2),
        'note': ParagraphStyle('Note', fontName=FONT, fontSize=8.5, leading=13,
                              alignment=TA_CENTER,
                              textColor=colors.HexColor('#666666'),
                              spaceBefore=2, spaceAfter=2),
    }


def runs_to_markup(runs):
    parts = []
    for run in runs:
        text = run.text
        if not text:
            continue
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        text = text.replace('\n', '<br/>')
        if run.font.bold:
            text = f'<b>{text}</b>'
        if run.font.italic:
            text = f'<i>{text}</i>'
        if run.font.color and run.font.color.rgb:
            c = str(run.font.color.rgb)
            if c != '000000':
                text = f'<font color="#{c}">{text}</font>'
        parts.append(text)
    return ''.join(parts)


def render_table(table, styles):
    nrows = len(table.rows)
    ncols = len(table.columns) if table.rows else 0
    if nrows == 0 or ncols == 0:
        return None
    data = []
    for i, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            cell_paras = []
            for p in cell.paragraphs:
                markup = runs_to_markup(p.runs)
                if markup.strip():
                    cell_paras.append(markup)
            cell_text = '<br/>'.join(cell_paras) if cell_paras else ''
            if i == 0 and cell_text:
                cell_text = f'<b>{cell_text}</b>'
            style = ParagraphStyle(
                f'Cell_{i}', fontName=FONT, fontSize=8.5, leading=12,
                textColor=colors.HexColor('#1f2328'),
            )
            row_data.append(Paragraph(cell_text, style))
        data.append(row_data)

    page_w = PAGE_SIZE[0] - 3 * cm
    col_widths = [page_w / ncols] * ncols
    if ncols == 2:
        col_widths = [page_w * 0.30, page_w * 0.70]
    elif ncols == 3:
        col_widths = [page_w * 0.18, page_w * 0.52, page_w * 0.30]
    elif ncols == 4:
        col_widths = [page_w * 0.10, page_w * 0.50, page_w * 0.18, page_w * 0.22]
    elif ncols == 5:
        col_widths = [page_w * 0.06, page_w * 0.12, page_w * 0.40, page_w * 0.16, page_w * 0.26]

    t = Table(data, colWidths=col_widths, repeatRows=1)
    style_cmds = [
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#888888')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('BACKGROUND', (0, 0), (0, 0), colors.HexColor('#D9E2F3')),  # header row first col
    ]
    # apply header bg to whole first row
    style_cmds[6] = ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9E2F3'))
    t.setStyle(TableStyle(style_cmds))
    return t


def docx_to_flowables():
    doc = Document(SRC)
    styles = make_styles()
    flow = []

    body = doc.element.body
    paragraphs = doc.paragraphs
    tables = doc.tables
    para_idx = 0
    table_idx = 0

    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            if para_idx < len(paragraphs):
                render_paragraph(paragraphs[para_idx], flow, styles)
                para_idx += 1
        elif tag == 'tbl':
            if table_idx < len(tables):
                t = render_table(tables[table_idx], styles)
                if t:
                    flow.append(t)
                    flow.append(Spacer(1, 4))
                table_idx += 1
    return flow


def render_paragraph(p, flow, styles):
    text = p.text.strip()
    if not text:
        # Empty paragraph but may contain image
        for run in p.runs:
            for blip in run._element.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId:
                    try:
                        part = p.part.related_parts[rId]
                        img_data = part.blob
                        ext = (part.partname.split('.')[-1] or 'png').lower()
                        tmp = Path(r'E:\claude\smart-workshop-erp\docs\_ops_inline.' + ext)
                        tmp.write_bytes(img_data)
                        from reportlab.platypus import Image as RLImage
                        img = RLImage(str(tmp), width=22 * cm, height=12 * cm)
                        flow.append(img)
                        flow.append(Spacer(1, 4))
                        return
                    except Exception as e:
                        print('  [WARN] inline image failed:', e)
        return

    # Non-empty text - also check for images
    for run in p.runs:
        for blip in run._element.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId:
                try:
                    part = p.part.related_parts[rId]
                    img_data = part.blob
                    ext = (part.partname.split('.')[-1] or 'png').lower()
                    tmp = Path(r'E:\claude\smart-workshop-erp\docs\_ops_inline.' + ext)
                    tmp.write_bytes(img_data)
                    from reportlab.platypus import Image as RLImage
                    img = RLImage(str(tmp), width=22 * cm, height=12 * cm)
                    flow.append(img)
                    flow.append(Spacer(1, 4))
                    return
                except Exception as e:
                    print('  [WARN] inline image failed:', e)

    if text.startswith('>') or text.startswith('「'):
        flow.append(Paragraph(text, styles['quote']))
        return

    if text in ('一、你的 3 步操作（记住这 3 步就够了）',
                '二、APP 界面速览（操作流程）',
                '三、5 类码（你必须分清）',
                '四、9 类异常处理（遇到不要慌，先按这个表）',
                '五、3 个红线（违反罚款 / 重罚）',
                '六、5 个"省自己时间"的小贴士',
                '七、找谁（不要自己扛）',
                '八、下料流程全景（泳道图）'):
        flow.append(Paragraph(text, styles['h1']))
        return

    if text.startswith('确认签收'):
        flow.append(Paragraph(text, styles['h2']))
        return

    if text.startswith('● '):
        b = Paragraph(runs_to_markup(p.runs), styles['body'])
        flow.append(b)
        return
    if text.startswith('🚫') or text.startswith('步骤') or text.startswith('口诀'):
        flow.append(Paragraph(runs_to_markup(p.runs), styles['body']))
        return

    # default body
    markup = runs_to_markup(p.runs)
    if markup.strip():
        flow.append(Paragraph(markup, styles['body']))


def make_canvas(total):
    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setFont(FONT, 8)
        canvas.setFillColor(colors.HexColor('#888888'))
        canvas.drawCentredString(
            PAGE_SIZE[0] / 2, PAGE_SIZE[1] - 0.8 * cm,
            '下料工操作手册（昆山佰泰胜专属 ERP V1.0）· A4 横向'
        )
        canvas.drawCentredString(
            PAGE_SIZE[0] / 2, 0.6 * cm,
            f'第 {doc.page} 页 / 共 {total} 页'
        )
        canvas.restoreState()
    return on_page


def main():
    print(f'[INFO] Reading: {SRC}')

    flow = docx_to_flowables()
    print(f'[INFO] {len(flow)} flowables')

    # First pass: count
    tmp_path = DST.with_suffix('.tmp.pdf')
    doc_tmp = SimpleDocTemplate(
        str(tmp_path), pagesize=PAGE_SIZE,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=1.2 * cm, bottomMargin=1.2 * cm,
    )
    counter = {'page': 0}
    def counting(canvas, doc):
        counter['page'] = max(counter['page'], doc.page)
    flow_for_count = [copy.deepcopy(f) for f in flow]
    doc_tmp.build(flow_for_count, onFirstPage=counting, onLaterPages=counting)
    total = max(counter['page'], 1)
    print(f'[INFO] Total pages: {total}')
    try:
        tmp_path.unlink()
    except FileNotFoundError:
        pass

    # Second pass
    flow_for_real = [copy.deepcopy(f) for f in flow]
    doc = SimpleDocTemplate(
        str(DST), pagesize=PAGE_SIZE,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=1.2 * cm, bottomMargin=1.2 * cm,
        title='下料工操作手册',
        author='河南晓评信息科技有限公司',
    )
    doc.build(flow_for_real, onFirstPage=make_canvas(total), onLaterPages=make_canvas(total))

    size = DST.stat().st_size
    print(f'[OK] PDF: {DST}')
    print(f'     Size: {size:,} bytes')
    print(f'     Pages: {total}')


if __name__ == '__main__':
    main()
