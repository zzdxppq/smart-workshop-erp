"""Convert contract-cnc-erp.md to a formal Word document (.docx)."""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"E:\claude\smart-workshop-erp\docs\contract-cnc-erp.md")
DST = Path(r"E:\claude\smart-workshop-erp\docs\contract-cnc-erp.docx")


def set_cell_bg(cell, color_hex):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell):
    """Add borders to a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_table(doc, rows, header_bg='D9E2F3'):
    """Create a table from list of lists. First row is header."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            add_runs_with_bold(p, str(cell_text), font_name='宋体', size=10.5)
            if i == 0:
                # Make all runs in header bold
                for run in p.runs:
                    run.font.bold = True
                set_cell_bg(cell, header_bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
    return table


def set_para_font(para, font_name='宋体', size=10.5, bold=False, color=None):
    """Apply Chinese font + size to paragraph."""
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)


def add_heading(doc, text, level=1):
    """Add a styled heading."""
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(24)
        run.font.bold = True
        set_para_font(p, '黑体', 24, True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
    elif level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        set_para_font(p, '黑体', 16, True)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        set_para_font(p, '黑体', 14, True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        set_para_font(p, '黑体', 12, True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)


def add_para(doc, text, indent=True, bold=False, align=None):
    """Add a normal paragraph with inline bold support."""
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2 字符
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    add_runs_with_bold(p, text, font_name='宋体', size=10.5, bold=bold)
    return p


def add_blockquote(doc, text, style='note'):
    """Add a styled paragraph for notes/callouts."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.right_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if style == 'note':
        add_runs_with_bold(p, '【注】' + text, font_name='楷体', size=10, color='666666')
    elif style == 'warn':
        add_runs_with_bold(p, '⚠ ' + text, font_name='楷体', size=10, color='A04040')
    return p


def add_runs_with_bold(para, text, font_name='宋体', size=10.5, bold=False, color=None):
    """Parse **text** as bold runs and add them to a paragraph."""
    # Split by **...** while keeping delimiters
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = para.add_run(inner)
            run.font.bold = True
        else:
            run = para.add_run(part)
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = run.font.bold or bold
        if color and not run.font.bold:
            run.font.color.rgb = RGBColor.from_string(color)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)


def add_hr(doc):
    """Add a horizontal rule (using a bottom-bordered empty paragraph)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '888888')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    doc.add_page_break()


def parse_md_to_docx():
    doc = Document()

    # Set default styles
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        # Header / footer
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
        # Header content
        header = section.header
        hp = header.paragraphs[0]
        hp.text = '昆山佰泰胜专属 ERP 系统定制开发项目 · 技术开发（委托）合同'
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            run.font.size = Pt(9)
            run.font.name = '宋体'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), '宋体')
        # Footer with page numbers
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run('第 ')
        run.font.size = Pt(9)
        run.font.name = '宋体'
        # Insert PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run = fp.add_run(' 页 / 共 ')
        run.font.size = Pt(9)
        run.font.name = '宋体'
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'NUMPAGES'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run = fp.add_run(' 页')
        run.font.size = Pt(9)
        run.font.name = '宋体'

    # COVER
    add_heading(doc, '技术开发（委托）合同', level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('合同编号：XP-ZPF202606082405')
    set_para_font(p, '宋体', 11, False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('签订日期：2026 年 6 月 ___ 日 · 签订地点：江苏省昆山市')
    set_para_font(p, '宋体', 11, False)
    add_hr(doc)
    add_para(doc, '本《技术开发（委托）合同》由以下双方签订：', indent=False, align='center')
    add_para(doc, '委托方（以下简称"甲方"）：昆山佰泰胜精密机械有限公司', indent=False, align='center')
    add_para(doc, '受托方（以下简称"乙方"）：河南晓评信息科技有限公司', indent=False, align='center')
    add_page_break(doc)

    # 首页：合同主体信息
    add_heading(doc, '一、合同主体信息', level=1)

    add_heading(doc, '（一）委托方（甲方）', level=3)
    add_table(doc, [
        ['项目', '内容'],
        ['单位名称', '昆山佰泰胜精密机械有限公司'],
        ['统一社会信用代码', '91320583MA21CG4M6J'],
        ['地址', '江苏省昆山市玉山镇成功路 155 号 3 号房'],
        ['联系电话', '0512-55106663'],
        ['联系人', '黄梓昀（151-0595-0281）'],
    ])

    add_heading(doc, '（二）受托方（乙方）', level=3)
    add_table(doc, [
        ['项目', '内容'],
        ['单位名称', '河南晓评信息科技有限公司'],
        ['统一社会信用代码', '91410100MA9EYQH00W'],
        ['法定代表人', '武蕴'],
        ['地址', '河南省郑州市高新技术产业开发区长椿路 11 号河南省国家大学科技园 1 号楼 3A 层 3A06'],
        ['联系人', '潘强（158-3710-7264）'],
    ])

    add_heading(doc, '（三）项目信息', level=3)
    add_table(doc, [
        ['项目', '内容'],
        ['项目名称', '昆山佰泰胜专属 ERP 系统定制开发项目'],
        ['合同金额', '人民币 ¥68,000.00 元（大写：陆万捌仟元整）'],
        ['含税类型', '6% 增值税专用发票（一般纳税人）'],
        ['货币单位', '人民币'],
        ['金额性质', '固定总价（除第 8.2 条约定变更外）'],
    ])

    add_heading(doc, '（四）甲方开票资料', level=3)
    add_table(doc, [
        ['字段', '内容'],
        ['名称', '昆山佰泰胜精密机械有限公司'],
        ['税号', '91320583MA21CG4M6J'],
        ['地址', '江苏省昆山市玉山镇成功路 155 号 3 号房'],
        ['电话', '0512-55106663'],
        ['开户银行', '江苏昆山农村商业银行股份有限公司城北支行'],
        ['行号', '314305206650'],
        ['账号', '3052250012012000010880'],
        ['开户行地址', '江苏省苏州市昆山市玉山镇北门路 609 号'],
    ])

    add_heading(doc, '（五）乙方收款账户', level=3)
    add_table(doc, [
        ['字段', '内容'],
        ['户名', '河南晓评信息科技有限公司'],
        ['开户银行', '中国工商银行郑州银河支行'],
        ['账号', '1702520719200056910'],
    ])

    add_page_break(doc)

    # 正文
    sections = parse_sections(SRC.read_text(encoding='utf-8'))
    for sec in sections:
        render_section(doc, sec)
        # Add some spacing between top-level sections
        if sec.get('level') == 1:
            doc.add_paragraph()

    # 签章页
    add_page_break(doc)

    add_para(doc, '本合同由双方在平等、自愿基础上签订，是双方真实意思表示。', indent=False, align='center')
    add_para(doc, '签订前请委托专业律师审阅，确保法律风险可控。', indent=False, align='center')
    add_para(doc, '签订后建议加盖骑缝章、双方各执一份、扫描件存档。', indent=False, align='center')
    doc.add_paragraph()
    doc.add_paragraph()

    add_heading(doc, '甲方（委托方）', level=2)
    add_table(doc, [
        ['项目', '内容'],
        ['单位名称', '昆山佰泰胜精密机械有限公司'],
        ['法定代表人 / 委托代理人', '（签字）：________________________'],
        ['统一社会信用代码', '91320583MA21CG4M6J'],
        ['公章', '（加盖）'],
        ['日期', '2026 年 6 月 ___ 日'],
    ])
    doc.add_paragraph()

    add_heading(doc, '乙方（受托方）', level=2)
    add_table(doc, [
        ['项目', '内容'],
        ['单位名称', '河南晓评信息科技有限公司'],
        ['法定代表人', '武蕴'],
        ['法定代表人 / 委托代理人', '（签字）：________________________'],
        ['统一社会信用代码', '91410100MA9EYQH00W'],
        ['公章', '（加盖）'],
        ['日期', '2026 年 6 月 ___ 日'],
    ])

    # Save to a temp path first, then atomically replace, to avoid 'Permission denied'
    # if the destination is being held open by Word/Reader.
    import os, time
    final_path = DST
    tmp_path = DST.with_suffix('.docx.tmp')
    doc.save(str(tmp_path))

    # If DST is locked, try rename-to-stale + write new copy
    replaced = False
    for attempt in range(3):
        try:
            if final_path.exists():
                # rename existing file to a stale name so the new file can take its slot
                stale = final_path.with_name(final_path.stem + f'.stale{int(time.time())}.docx')
                try:
                    final_path.rename(stale)
                except PermissionError:
                    time.sleep(0.5)
                    continue
                else:
                    try:
                        stale.unlink()
                    except OSError:
                        pass
            os.replace(str(tmp_path), str(final_path))
            replaced = True
            break
        except PermissionError:
            time.sleep(0.5)

    if not replaced:
        # Last resort: write to a timestamped sibling so the deliverable still gets produced
        alt = final_path.with_name(final_path.stem + f'.{int(time.time())}.docx')
        os.replace(str(tmp_path), str(alt))
        print(f"[WARN] Could not overwrite {final_path.name}, saved to {alt.name}")
        final_path = alt

    sys.stdout.reconfigure(encoding='utf-8')
    print(f"[OK] Contract Word: {final_path}")
    print(f"     Size: {final_path.stat().st_size:,} bytes")


# ================== MD Parsing ==================

def parse_sections(md_text):
    """Parse MD into a list of section dicts."""
    lines = md_text.split('\n')
    sections = []
    i = 0
    n = len(lines)
    # Find start: skip until we hit the first article (### 第一条) or 附件
    # Skip the cover/homepage part (合同首页 / 一、合同主体信息)
    # because the script already hardcodes the homepage as tables.
    started = False
    while i < n:
        line = lines[i]
        if not started:
            if re.match(r'^##\s+合同正文', line):
                started = True
                i += 1
                continue
            else:
                i += 1
                continue
        # Parse heading
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            # Collect content until next heading
            content = []
            i += 1
            while i < n:
                nxt = lines[i]
                if re.match(r'^#{1,4}\s+', nxt):
                    break
                content.append(nxt)
                i += 1
            sections.append({
                'level': level,
                'title': title,
                'content': '\n'.join(content).strip(),
            })
        else:
            i += 1
    return sections


def render_section(doc, sec):
    """Render a section to docx."""
    level = sec['level']
    title = sec['title']
    content = sec['content']

    # Articles (### 第N条) should be H2 (16pt) for visibility
    if re.match(r'^第[一二三四五六七八九十]+条', title):
        add_heading(doc, title, level=2)
        # Force a more prominent styling
        if doc.paragraphs:
            last_para = doc.paragraphs[-1]
            for run in last_para.runs:
                run.font.size = Pt(15)
                run.font.bold = True
                run.font.name = '黑体'
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), '黑体')
        # IMPORTANT: continue to render the section's body (子条款/表格/列表),
        # otherwise the content under every "第N条" would be silently dropped.
        # Fall through to the content renderer below.
    if level == 1:
        add_heading(doc, title, level=1)
    elif level == 2:
        add_heading(doc, title, level=2)
    elif level == 3:
        add_heading(doc, title, level=3)
    else:
        add_heading(doc, title, level=3)

    if not content:
        return

    # Parse content line by line
    lines = content.split('\n')
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()

        # Blockquote
        if line.startswith('>'):
            quote_lines = []
            while i < n and lines[i].rstrip().startswith('>'):
                quote_lines.append(re.sub(r'^>\s?', '', lines[i].rstrip()))
                i += 1
            quote_text = ' '.join(q.strip() for q in quote_lines if q.strip())
            if quote_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.74)
                p.paragraph_format.right_indent = Cm(0.74)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if quote_text.startswith('【注】') or quote_text.startswith('**注**'):
                    add_runs_with_bold(p, '【注】' + quote_text.replace('**注**', '').replace('【注】', ''),
                                       font_name='楷体', size=10, color='666666')
                elif quote_text.startswith('⚠') or quote_text.startswith('**⚠'):
                    add_runs_with_bold(p, '⚠ ' + quote_text.replace('**⚠**', '').replace('⚠', ''),
                                       font_name='楷体', size=10, color='A04040')
                else:
                    add_runs_with_bold(p, '「' + quote_text + '」', font_name='楷体', size=10)
            continue

        # Table
        if line.startswith('|'):
            table_rows = []
            while i < n and lines[i].rstrip().startswith('|'):
                row_text = lines[i].rstrip()
                # Skip separator line |---|---|
                if re.match(r'^\|[\s\-|:]+\|$', row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
            if table_rows:
                add_table(doc, table_rows)
                doc.add_paragraph()
            continue

        # Horizontal rule
        if line.strip() == '---':
            add_hr(doc)
            i += 1
            continue

        # Heading within content (sub-sub-section)
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            sub_level = len(m.group(1))
            sub_title = m.group(2).strip()
            add_heading(doc, sub_title, level=3 if sub_level <= 3 else 3)
            i += 1
            continue

        # Bold/italic paragraphs or plain text
        if line.strip() == '':
            i += 1
            continue

        # List item: process directly, do NOT go through plain text path
        m_ul_curr = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        m_ol_curr = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if m_ul_curr or m_ol_curr:
            # Fall through to list processor below
            pass
        else:
            # Plain text paragraph: accumulate subsequent non-list lines
            text = line.strip()
            para_lines = [text]
            i += 1
            while i < n:
                nxt = lines[i].rstrip()
                if (nxt == '' or nxt.startswith('#') or nxt.startswith('|') or
                        nxt.startswith('>') or nxt.strip() == '---' or
                        re.match(r'^\s*[-*+]\s+', nxt) or re.match(r'^\s*\d+\.\s+', nxt)):
                    break
                para_lines.append(nxt)
                i += 1
            p_text = ' '.join(l.strip() for l in para_lines if l.strip())
            if p_text:
                add_para(doc, p_text, indent=True)
            continue

        # Process list items (current line + any subsequent list items)
        while i < n:
            nxt = lines[i].rstrip()
            if not nxt:
                i += 1
                continue
            m_ul = re.match(r'^(\s*)[-*+]\s+(.+)$', nxt)
            m_ol = re.match(r'^(\s*)(\d+)\.\s+(.+)$', nxt)
            if m_ul:
                indent_spaces = len(m_ul.group(1))
                bullet_text = m_ul.group(2)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.74 + 0.5 * (indent_spaces // 2))
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                add_runs_with_bold(p, '● ' + bullet_text, font_name='宋体', size=10.5)
                i += 1
            elif m_ol:
                indent_spaces = len(m_ol.group(1))
                num = m_ol.group(2)
                list_text = m_ol.group(3)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.74 + 0.5 * (indent_spaces // 2))
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                add_runs_with_bold(p, f'{num}. ' + list_text, font_name='宋体', size=10.5)
                i += 1
            else:
                break


if __name__ == '__main__':
    parse_md_to_docx()
